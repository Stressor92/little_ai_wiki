"""
converters/docx_conv.py
DOCX/ODT → Markdown via python-docx

Verarbeitet:
  - Paragraph-Styles (Heading 1/2/3, Normal, List, Quote)
  - Inline-Formatierung (Bold, Italic, Underline, Code)
  - Tabellen (vollständig mit Header-Erkennung)
  - Nummerierte und nicht-nummerierte Listen
  - Seitenumbrüche als Kapitel-Trenner
"""

from __future__ import annotations
import logging
import re
import os
import subprocess
import tempfile
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path

from .base import (
    BaseConverter, Block, BlockType, Chapter,
    Document, DocumentMeta,
)

log = logging.getLogger(__name__)

# python-docx Style-Namen → Block-Typen
STYLE_MAP: dict[str, BlockType] = {
    "heading 1":              BlockType.HEADING_1,
    "heading 2":              BlockType.HEADING_2,
    "heading 3":              BlockType.HEADING_3,
    "heading 4":              BlockType.HEADING_4,
    "title":                  BlockType.HEADING_1,
    "subtitle":               BlockType.HEADING_2,
    "block text":             BlockType.QUOTE,
    "quote":                  BlockType.QUOTE,
    "intense quote":          BlockType.QUOTE,
    "caption":                BlockType.CAPTION,
    "list paragraph":         BlockType.LIST_ITEM,
    "list bullet":            BlockType.LIST_ITEM,
    "list number":            BlockType.LIST_ITEM,
    # Deutsche Style-Namen
    "überschrift 1":          BlockType.HEADING_1,
    "überschrift 2":          BlockType.HEADING_2,
    "überschrift 3":          BlockType.HEADING_3,
    "überschrift 4":          BlockType.HEADING_4,
    "zitat":                  BlockType.QUOTE,
    "listenabsatz":           BlockType.LIST_ITEM,
}


class DOCXConverter(BaseConverter):
    SUPPORTED_EXTENSIONS = (".docx", ".docm", ".doc", ".odt")

    def convert(self, input_path: Path) -> Document:
        source_path = input_path
        temp_dir: tempfile.TemporaryDirectory[str] | None = None

        if input_path.suffix.lower() == ".odt" and self._find_soffice() is None:
            return self._convert_odt_direct(input_path)

        if input_path.suffix.lower() == ".doc" and self._find_soffice() is None:
            source_path, temp_dir = self._convert_doc_with_word(input_path)
        elif input_path.suffix.lower() in (".doc", ".odt"):
            temp_dir = tempfile.TemporaryDirectory()
            source_path = self._convert_to_docx_with_soffice(input_path, Path(temp_dir.name))

        try:
            from docx import Document as DocxDocument
            from docx.oxml.ns import qn
        except ImportError as exc:
            if temp_dir is not None:
                temp_dir.cleanup()
            raise ImportError("python-docx fehlt: pip install python-docx") from exc

        try:
            docx     = DocxDocument(str(source_path))
            meta     = self._extract_meta(docx, input_path)
            doc      = Document(meta=meta)
            chapter  = doc.add_chapter(meta.title or input_path.stem)

            list_counter: dict[int, int] = {}   # Nummerierung geordneter Listen

            for element in docx.element.body:
                tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

                if tag == "p":
                    from docx.text.paragraph import Paragraph
                    para = Paragraph(element, docx)
                    block = self._process_paragraph(para, list_counter)
                    if block is None:
                        continue

                    if block.type == BlockType.HEADING_1:
                        chapter = doc.add_chapter(block.text)
                        doc.meta.chapters = len(doc.chapters)
                    else:
                        chapter.blocks.append(block)

                elif tag == "tbl":
                    from docx.table import Table
                    table = Table(element, docx)
                    table_blocks = self._process_table(table)
                    chapter.blocks.extend(table_blocks)

                elif tag == "sectPr":
                    # Seitenumbruch / neue Section → bleibt im selben Kapitel
                    pass

            return doc
        finally:
            if temp_dir is not None:
                temp_dir.cleanup()

    def _convert_to_docx_with_soffice(self, input_path: Path, out_dir: Path) -> Path:
        cmd = self._find_soffice()
        if cmd is None:
            raise RuntimeError(
                f"{input_path.suffix.lower()} requires LibreOffice (soffice) in PATH for conversion to .docx"
            )

        result = subprocess.run(
            [
                cmd,
                "--headless",
                "--convert-to",
                "docx",
                "--outdir",
                str(out_dir),
                str(input_path),
            ],
            capture_output=True,
            text=True,
            timeout=120,
        )

        if result.returncode != 0:
            raise RuntimeError(f"LibreOffice conversion failed: {result.stderr[:500]}")

        converted = out_dir / f"{input_path.stem}.docx"
        if not converted.exists():
            raise RuntimeError("LibreOffice did not produce .docx output.")
        return converted

    @staticmethod
    def _find_soffice() -> str | None:
        for candidate in ("soffice", "libreoffice"):
            try:
                check = subprocess.run(
                    [candidate, "--version"],
                    capture_output=True,
                    text=True,
                    timeout=5,
                )
                if check.returncode == 0:
                    return candidate
            except (FileNotFoundError, subprocess.TimeoutExpired):
                continue
        return None

    def _convert_doc_with_word(self, input_path: Path) -> tuple[Path, tempfile.TemporaryDirectory[str]]:
        """Convert legacy .doc to .docx using Microsoft Word COM on Windows."""
        if os.name != "nt":
            raise RuntimeError(".doc conversion without LibreOffice is only supported on Windows.")

        try:
            import pythoncom
            import win32com.client
        except ImportError as exc:
            raise RuntimeError(
                ".doc fallback requires pywin32 and Microsoft Word (pip install pywin32)."
            ) from exc

        temp_dir = tempfile.TemporaryDirectory()
        out_docx = Path(temp_dir.name) / f"{input_path.stem}.docx"

        pythoncom.CoInitialize()
        word = None
        doc = None
        try:
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = 0

            doc = word.Documents.Open(str(input_path.resolve()))
            # 16 = wdFormatDocumentDefault (.docx)
            doc.SaveAs(str(out_docx), FileFormat=16)
            doc.Close(False)
            doc = None

            if not out_docx.exists():
                raise RuntimeError("Microsoft Word conversion did not produce .docx output.")
            return out_docx, temp_dir
        except Exception as exc:
            raise RuntimeError(
                ".doc conversion failed. Ensure Microsoft Word is installed and file is not locked."
            ) from exc
        finally:
            if doc is not None:
                try:
                    doc.Close(False)
                except Exception:
                    pass
            if word is not None:
                try:
                    word.Quit()
                except Exception:
                    pass
            pythoncom.CoUninitialize()

    def _convert_odt_direct(self, input_path: Path) -> Document:
        meta = self._make_meta(input_path, "ODT")
        doc = Document(meta=meta)
        chapter = doc.add_chapter(meta.title or input_path.stem)

        text_ns = "urn:oasis:names:tc:opendocument:xmlns:text:1.0"
        office_ns = "urn:oasis:names:tc:opendocument:xmlns:office:1.0"
        ns = {"text": text_ns, "office": office_ns}

        with zipfile.ZipFile(input_path, "r") as zf:
            if "content.xml" not in zf.namelist():
                raise RuntimeError("Invalid ODT: content.xml missing")
            content_xml = zf.read("content.xml")

        root = ET.fromstring(content_xml)
        body = root.find(".//office:body/office:text", ns)
        if body is None:
            return doc

        for elem in body.iter():
            if not isinstance(elem.tag, str):
                continue

            tag = elem.tag
            text = " ".join(part.strip() for part in elem.itertext() if part and part.strip())
            if not text:
                continue

            if tag == f"{{{text_ns}}}h":
                level_raw = elem.attrib.get(f"{{{text_ns}}}outline-level", "1")
                try:
                    level = int(level_raw)
                except ValueError:
                    level = 1

                if level <= 1:
                    chapter = doc.add_chapter(text)
                    doc.meta.chapters = len(doc.chapters)
                elif level == 2:
                    chapter.blocks.append(Block(type=BlockType.HEADING_2, text=text))
                elif level == 3:
                    chapter.blocks.append(Block(type=BlockType.HEADING_3, text=text))
                else:
                    chapter.blocks.append(Block(type=BlockType.HEADING_4, text=text))
            elif tag == f"{{{text_ns}}}p":
                chapter.blocks.append(Block(type=BlockType.PARAGRAPH, text=text))

        doc.meta.chapters = len(doc.chapters)
        return doc

    # ── Metadaten ─────────────────────────────────────────────────────────────

    def _extract_meta(self, docx, path: Path) -> DocumentMeta:
        cp = docx.core_properties
        description = getattr(cp, "description", "") or getattr(cp, "comments", "") or ""
        return DocumentMeta(
            title       = cp.title    or path.stem,
            author      = cp.author   or cp.last_modified_by or "",
            date        = str(cp.created)[:10] if cp.created else "",
            description = description or cp.subject or "",
            source_file = path.name,
            source_fmt  = "DOCX",
        )

    # ── Paragraphen ───────────────────────────────────────────────────────────

    def _process_paragraph(self, para, list_counter: dict) -> Block | None:
        text = self._extract_inline_text(para)
        if not text:
            return None

        style_name = para.style.name.lower() if para.style else ""

        # Style-Mapping
        if style_name in STYLE_MAP:
            btype = STYLE_MAP[style_name]
        # Outline-Level aus Paragraph-Format
        elif hasattr(para._p, 'pPr') and para._p.pPr is not None:
            outline = para._p.pPr.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl"
            )
            if outline is not None:
                lvl = int(outline.get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", 9
                ))
                level_map = {0: BlockType.HEADING_1, 1: BlockType.HEADING_2,
                             2: BlockType.HEADING_3, 3: BlockType.HEADING_4}
                btype = level_map.get(lvl, BlockType.PARAGRAPH)
            else:
                btype = BlockType.PARAGRAPH
        else:
            btype = BlockType.PARAGRAPH

        # Listen-Typ aus numId / ilvl
        num_id = self._get_num_id(para)
        ilvl   = self._get_ilvl(para)
        is_ordered = self._is_ordered_list(para)

        if num_id is not None or btype == BlockType.LIST_ITEM:
            if is_ordered:
                cnt_key = (num_id or 0, ilvl)
                list_counter[cnt_key] = list_counter.get(cnt_key, 0) + 1
                ordered = list_counter[cnt_key]
            else:
                ordered = False
            return Block(
                type=BlockType.LIST_ITEM,
                text=text,
                level=ilvl,
                ordered=ordered,
            )

        return Block(type=btype, text=text)

    def _extract_inline_text(self, para) -> str:
        """Extrahiert Text mit Inline-Markdown-Formatierung."""
        from docx.oxml.ns import qn
        parts = []

        for run in para.runs:
            if not run.text:
                continue
            t = run.text
            # Code: Courier/Mono-Font → `code`
            font_name = (run.font.name or "").lower()
            if any(m in font_name for m in ("courier", "mono", "consolas", "code")):
                t = f"`{t}`"
            elif run.bold and run.italic:
                t = f"***{t}***"
            elif run.bold:
                t = f"**{t}**"
            elif run.italic:
                t = f"*{t}*"
            parts.append(t)

        return "".join(parts).strip()

    # ── Tabellen ──────────────────────────────────────────────────────────────

    def _process_table(self, table) -> list[Block]:
        blocks = []
        if not table.rows:
            return blocks

        # Erste Zeile als Header
        header_row = table.rows[0]
        blocks.append(Block(
            type=BlockType.TABLE_HEADER,
            cells=[self._cell_text(c) for c in header_row.cells],
            text="",
        ))

        for row in table.rows[1:]:
            blocks.append(Block(
                type=BlockType.TABLE_ROW,
                cells=[self._cell_text(c) for c in row.cells],
                text="",
            ))

        return blocks

    @staticmethod
    def _cell_text(cell) -> str:
        texts = []
        for para in cell.paragraphs:
            t = " ".join(r.text for r in para.runs if r.text)
            if t.strip():
                texts.append(t.strip())
        return " ".join(texts)

    # ── Listen-Hilfsmethoden ──────────────────────────────────────────────────

    @staticmethod
    def _get_num_id(para) -> int | None:
        try:
            pPr = para._p.pPr
            if pPr is None:
                return None
            numPr = pPr.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
            )
            if numPr is None:
                return None
            numId = numPr.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numId"
            )
            return int(numId.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", 0
            )) if numId is not None else None
        except Exception:
            return None

    @staticmethod
    def _get_ilvl(para) -> int:
        try:
            pPr = para._p.pPr
            if pPr is None:
                return 0
            numPr = pPr.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
            )
            if numPr is None:
                return 0
            ilvl = numPr.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ilvl"
            )
            return int(ilvl.get(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", 0
            )) if ilvl is not None else 0
        except Exception:
            return 0

    @staticmethod
    def _is_ordered_list(para) -> bool:
        style_name = (para.style.name or "").lower()
        return "number" in style_name or "numeri" in style_name
